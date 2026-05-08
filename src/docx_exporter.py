from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


FONT_NAME = "Times New Roman"
BODY_FONT_SIZE = 13


@dataclass(frozen=True)
class ExportProfile:
    font_name: str = FONT_NAME
    body_font_size: int = BODY_FONT_SIZE
    line_spacing: float = 1.15
    paragraph_spacing_after_pt: int = 6
    first_line_indent_cm: float = 1.27
    top_margin_mm: int = 20
    bottom_margin_mm: int = 20
    left_margin_mm: int = 30
    right_margin_mm: int = 15


def export_draft_to_docx(
    *,
    draft: str,
    doc_type: str,
    agency_parent: str = "[TÊN CƠ QUAN CHỦ QUẢN]",
    agency_name: str = "[TÊN CƠ QUAN BAN HÀNH]",
    place_name: str = "...",
    profile: ExportProfile | None = None,
) -> bytes:
    profile = profile or ExportProfile()
    document = Document()
    _setup_document(document, profile)

    lines = _clean_lines(draft)
    date_line = _find_date_line(lines, place_name)
    body_lines = _remove_plain_header(lines)

    _add_header_table(
        document=document,
        agency_parent=agency_parent,
        agency_name=agency_name,
        profile=profile,
    )
    _add_date_line(document, date_line, profile)
    _add_body_lines(document, body_lines, doc_type, profile)
    _add_page_number(document, profile)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _setup_document(document: Document, profile: ExportProfile) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(profile.top_margin_mm)
    section.bottom_margin = Mm(profile.bottom_margin_mm)
    section.left_margin = Mm(profile.left_margin_mm)
    section.right_margin = Mm(profile.right_margin_mm)

    style = document.styles["Normal"]
    style.font.name = profile.font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), profile.font_name)
    style.font.size = Pt(profile.body_font_size)
    style.paragraph_format.line_spacing = profile.line_spacing
    style.paragraph_format.space_after = Pt(profile.paragraph_spacing_after_pt)


def _add_header_table(
    *,
    document: Document,
    agency_parent: str,
    agency_name: str,
    profile: ExportProfile,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    _remove_table_borders(table)

    left_cell, right_cell = table.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    _set_cell_width(left_cell, Cm(7.2))
    _set_cell_width(right_cell, Cm(8.5))

    _set_cell_text(
        left_cell,
        [
            (agency_parent.upper(), 12, False),
            (agency_name.upper(), 12, True),
            ("-" * 18, 12, False),
        ],
        WD_ALIGN_PARAGRAPH.CENTER,
        profile,
    )
    _set_cell_text(
        right_cell,
        [
            ("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", 13, True),
            ("Độc lập - Tự do - Hạnh phúc", 14, True),
            ("-" * 28, 12, False),
        ],
        WD_ALIGN_PARAGRAPH.CENTER,
        profile,
    )


def _add_date_line(
    document: Document,
    date_line: str,
    profile: ExportProfile,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(date_line)
    _style_run(run, profile.font_name, 13, italic=True)


def _add_body_lines(
    document: Document,
    lines: list[str],
    doc_type: str,
    profile: ExportProfile,
) -> None:
    for line in lines:
        if not line:
            document.add_paragraph()
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = profile.line_spacing
        paragraph.paragraph_format.space_after = Pt(profile.paragraph_spacing_after_pt)

        alignment = _line_alignment(line, doc_type)
        paragraph.alignment = alignment
        if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and not line.startswith("- "):
            paragraph.paragraph_format.first_line_indent = Cm(profile.first_line_indent_cm)

        run = paragraph.add_run(line)
        size, bold, italic = _line_style(line, doc_type)
        _style_run(run, profile.font_name, size, bold=bold, italic=italic)


def _add_page_number(document: Document, profile: ExportProfile) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True

    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    _style_run(run, profile.font_name, 13)

    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def _set_cell_text(
    cell,
    lines: list[tuple[str, int, bool]],
    alignment,
    profile: ExportProfile,
) -> None:
    cell.text = ""
    for index, (text, size, bold) in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        _style_run(run, profile.font_name, size, bold=bold)


def _style_run(
    run,
    font_name: str,
    size: int,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _line_alignment(line: str, doc_type: str):
    upper_line = line.upper()
    centered_exact = {
        "THÔNG BÁO",
        "TỜ TRÌNH",
        "QUYẾT ĐỊNH",
        "QUYẾT ĐỊNH:",
        "NGƯỜI KÝ",
        "NGƯỜI TRÌNH",
    }
    if upper_line in centered_exact or upper_line == doc_type.upper():
        return WD_ALIGN_PARAGRAPH.CENTER
    if line.startswith("V/v ") or line.startswith("Về "):
        return WD_ALIGN_PARAGRAPH.CENTER
    if line.startswith("Số:"):
        return WD_ALIGN_PARAGRAPH.LEFT
    if line.startswith("Nơi nhận:"):
        return WD_ALIGN_PARAGRAPH.LEFT
    if line.startswith("- "):
        return WD_ALIGN_PARAGRAPH.LEFT
    if line.startswith("[Chức vụ"):
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.JUSTIFY


def _line_style(line: str, doc_type: str) -> tuple[int, bool, bool]:
    upper_line = line.upper()
    if upper_line in {"THÔNG BÁO", "TỜ TRÌNH", "QUYẾT ĐỊNH", doc_type.upper()}:
        return 14, True, False
    if upper_line == "QUYẾT ĐỊNH:":
        return 13, True, False
    if line.startswith("V/v ") or line.startswith("Về "):
        return 14, True, False
    if line.startswith("Số:"):
        return 13, False, False
    if line.startswith("Kính gửi:"):
        return 13, True, False
    if line.startswith("Điều "):
        return 13, True, False
    if line.startswith("Nguồn tham khảo"):
        return 13, True, False
    if line.startswith("Nơi nhận:"):
        return 12, True, True
    if upper_line in {"NGƯỜI KÝ", "NGƯỜI TRÌNH"}:
        return 13, True, False
    if line.startswith("Ghi chú kiểm soát:"):
        return 12, False, True
    if _is_numbered_heading(line):
        return 13, True, False
    return BODY_FONT_SIZE, False, False


def _clean_lines(draft: str) -> list[str]:
    return [line.rstrip() for line in draft.replace("\r\n", "\n").split("\n")]


def _find_date_line(lines: list[str], place_name: str) -> str:
    for line in lines:
        if "ngày" in line.lower() and "tháng" in line.lower() and "năm" in line.lower():
            if line.strip().startswith("...,") and place_name != "...":
                return line.replace("...,", f"{place_name},", 1)
            return line.strip()

    return f"{place_name}, ngày ... tháng ... năm ..."


def _remove_plain_header(lines: list[str]) -> list[str]:
    removable_prefixes = (
        "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM",
        "Độc lập - Tự do - Hạnh phúc",
    )
    cleaned: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        if stripped in removable_prefixes:
            continue
        if "ngày" in stripped.lower() and "tháng" in stripped.lower() and "năm" in stripped.lower():
            continue
        cleaned.append(stripped)

    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    return cleaned


def _is_numbered_heading(line: str) -> bool:
    return bool(line[:2].strip(".").isdigit()) and "." in line[:4]


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tbl_pr.append(borders)


def _set_cell_width(cell, width) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.cm * 567)))
    tc_w.set(qn("w:type"), "dxa")
