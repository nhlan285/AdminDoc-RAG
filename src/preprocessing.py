from __future__ import annotations
import hashlib
import json
import re
import io
import docx
from pypdf import PdfReader
from pathlib import Path
from typing import Any
import unicodedata
from src.documents import Document

DEFAULT_CHUNK_SIZE_WORDS = 120
DEFAULT_CHUNK_OVERLAP_WORDS = 20

ND30_VALID_DOC_TYPES = {
    "Nghị quyết", "Quyết định", "Chỉ thị", "Quy chế", "Quy định", "Hướng dẫn",
    "Thông cáo", "Thông báo", "Báo cáo", "Biên bản", "Tờ trình",
    "Chương trình", "Kế hoạch", "Phương án", "Đề án", "Dự án",
    "Công văn", "Công điện", "Bản ghi nhớ", "Bản thỏa thuận", "Hợp đồng",
    "Giấy ủy quyền", "Giấy mời", "Giấy giới thiệu", "Giấy nghỉ phép",
    "Phiếu gửi", "Phiếu chuyển", "Phiếu báo", "Thư công"
}

def clean_text(text: str) -> str:
    text = text.replace("\ufeff", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    cleaned_lines: list[str] = []
    previous_blank = False
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            if not previous_blank:
                cleaned_lines.append("")
            previous_blank = True
            continue
        cleaned_lines.append(line)
        previous_blank = False
    return "\n".join(cleaned_lines).strip()

def normalize_doc_type(raw_type: str) -> str:
    clean_type = str(raw_type).strip()
    if clean_type in ND30_VALID_DOC_TYPES:
        return clean_type
    lower_type = clean_type.lower()
    for valid_type in ND30_VALID_DOC_TYPES:
        if valid_type.lower() in lower_type:
            return valid_type
    return "Công văn"
def auto_detect_doc_type(filename: str, content: str) -> str:
    """Tự động phân loại văn bản hỗ trợ cả tên file không dấu và có dấu."""
    lower_filename = filename.lower()
    
    # 1. Bộ từ điển ánh xạ file không dấu (Dành cho file mẫu)
    filename_map = {
        "nghi_quyet": "Nghị quyết", "quyet_dinh": "Quyết định", "chi_thi": "Chỉ thị",
        "quy_che": "Quy chế", "quy_dinh": "Quy định", "huong_dan": "Hướng dẫn",
        "thong_cao": "Thông cáo", "thong_bao": "Thông báo", "bao_cao": "Báo cáo",
        "bien_ban": "Biên bản", "to_trinh": "Tờ trình", "chuong_trinh": "Chương trình",
        "ke_hoach": "Kế hoạch", "phuong_an": "Phương án", "de_an": "Đề án",
        "du_an": "Dự án", "cong_van": "Công văn", "cong_dien": "Công điện",
        "ghi_nho": "Bản ghi nhớ", "thoa_thuan": "Bản thỏa thuận", "hop_dong": "Hợp đồng",
        "uy_quyen": "Giấy ủy quyền", "giay_moi": "Giấy mời", "gioi_thieu": "Giấy giới thiệu",
        "nghi_phep": "Giấy nghỉ phép", "phieu_gui": "Phiếu gửi", "phieu_chuyen": "Phiếu chuyển",
        "phieu_bao": "Phiếu báo", "thu_cong": "Thư công"
    }
    
    # Quét theo từ khóa không dấu trước
    for key, valid_type in filename_map.items():
        if key in lower_filename:
            return valid_type

    # 2. Quét theo tên file có dấu chuẩn
    for valid_type in ND30_VALID_DOC_TYPES:
        if valid_type.lower() in lower_filename:
            return valid_type
            
    # 3. Quét 1000 ký tự đầu của nội dung nếu tên file không có thông tin
    lower_content = content[:1000].lower()
    for valid_type in ND30_VALID_DOC_TYPES:
        if valid_type.lower() in lower_content:
            return valid_type
            
    return "Công văn"

def extract_text_from_binary(file_bytes: bytes, filename: str) -> str:
    extension = filename.split('.')[-1].lower()
    try:
        if extension == 'docx':
            doc = docx.Document(io.BytesIO(file_bytes))
            return "\n".join([para.text for para in doc.paragraphs])
        elif extension == 'pdf':
            reader = PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text
        else:
            return file_bytes.decode("utf-8-sig", errors="replace")
    except Exception as e:
        raise ValueError(f"Không thể đọc file {filename}: {str(e)}")

def parse_documents_from_text(
    *,
    filename: str,
    text: str,
    default_doc_type: str,
    source_kind: str = "system",
) -> list[Document]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".json":
        try:
            return _parse_json_documents(
                text=text,
                filename=filename,
                default_doc_type=default_doc_type,
                source_kind=source_kind,
            )
        except Exception:
            pass

    cleaned = clean_text(text)
    if not cleaned:
        return []
    title = _guess_title(filename, cleaned)
    
    if _is_auto_doc_type(default_doc_type):
        doc_type = auto_detect_doc_type(filename, cleaned)
    else:
        doc_type = normalize_doc_type(default_doc_type)

    return [
        Document(
            id=_stable_id("DOC", filename, title, cleaned),
            title=title,
            doc_type=doc_type,
            source=filename,
            content=cleaned,
            source_kind=source_kind,
        )
    ]

def build_chunks(
    documents: list[Document],
    *,
    chunk_size_words: int = DEFAULT_CHUNK_SIZE_WORDS,
    overlap_words: int = DEFAULT_CHUNK_OVERLAP_WORDS,
) -> list[Document]:
    chunks: list[Document] = []
    for document in documents:
        cleaned_content = clean_text(document.content)
        words = cleaned_content.split()
        if not words:
            continue
        c_size = max(40, chunk_size_words)
        o_words = max(0, min(overlap_words, c_size // 2))
        start = 0
        doc_chunks_text: list[str] = []
        while start < len(words):
            end = min(start + c_size, len(words))
            doc_chunks_text.append(" ".join(words[start:end]))
            if end == len(words):
                break
            start = end - o_words
        total = len(doc_chunks_text)
        for index, chunk_text in enumerate(doc_chunks_text, start=1):
            chunks.append(
                Document(
                    id=f"{document.id}-CH{index:03d}",
                    title=document.title,
                    doc_type=document.doc_type,
                    source=document.source,
                    content=chunk_text,
                    source_kind=document.source_kind,
                    parent_id=document.id,
                    chunk_index=index,
                    total_chunks=total,
                )
            )
    return chunks

def _parse_json_documents(
    *,
    text: str,
    filename: str,
    default_doc_type: str,
    source_kind: str,
) -> list[Document]:
    raw_data = json.loads(text)
    raw_documents = raw_data.get("documents", []) if isinstance(raw_data, dict) else raw_data
    documents: list[Document] = []
    for index, raw_item in enumerate(raw_documents, start=1):
        if not isinstance(raw_item, dict):
            continue
        content = clean_text(str(raw_item.get("content") or raw_item.get("text") or ""))
        if not content:
            continue
        title = str(raw_item.get("title") or _guess_title(filename, content)).strip()
        source = str(raw_item.get("source") or filename).strip()
        doc_type = normalize_doc_type(raw_item.get("doc_type") or default_doc_type)
        document_id = str(raw_item.get("id") or "").strip()
        if not document_id:
            document_id = _stable_id("DOC", filename, title, str(index), content)
        documents.append(
            Document(
                id=document_id,
                title=title,
                doc_type=doc_type,
                source=source,
                content=content,
                source_kind=source_kind,
            )
        )
    return documents

def _guess_title(filename: str, content: str) -> str:
    for line in content.splitlines():
        line = line.strip(" #\t")
        if line:
            return line[:90]
    return Path(filename).stem.replace("_", " ").strip() or "Tài liệu chưa đặt tên"

def _stable_id(prefix: str, *parts: Any) -> str:
    raw_value = "|".join(str(part) for part in parts)
    digest = hashlib.sha1(raw_value.encode("utf-8")).hexdigest()[:10].upper()
    return f"{prefix}-{digest}"

def _is_auto_doc_type(raw_type: str) -> bool:
    normalized = str(raw_type).strip().lower().replace("đ", "d")
    normalized = unicodedata.normalize("NFD", normalized)
    normalized = "".join(
        char for char in normalized if unicodedata.category(char) != "Mn"
    )
    return normalized in {"auto", "automatic"} or "tu dong" in normalized
